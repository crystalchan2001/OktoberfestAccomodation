from docx import Document
from docx.shared import Pt
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.shared import RGBColor
from docx.text.run import Run
import os

class Writer():
    def __init__(self, listings, path, date):
        self.listings = listings
        self.path = path
        self.date = date

        self.doc = Document()
        self.setStyle()
        self.write()
        self.save()
        self.open()

    def setStyle(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)

    def write(self):
        self.doc.add_paragraph(f"{len(self.listings)} results from {self.date}")
        idx = 1
        for listing in self.listings:
            screenshot = listing.getScreenshot()
            commute = listing.getCommute()
            pricePer = listing.getPricePer()
            total = listing.getTotal()
            url = listing.getUrl()

            self.doc.add_picture(screenshot)
            p = self.doc.add_paragraph('Link to')
            self.addHyperlink(p, 'listing', url)

            self.doc.add_paragraph(commute)
            self.doc.add_paragraph(pricePer)
            self.doc.add_paragraph(total)

    def addHyperlink(paragraph, text, url):
        part = paragraph.part
        r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:r'), r_id)

        new_run = Run(OxmlElement('w:r'), paragraph)
        new_run.text = text
        new_run.font.color.rgb = RGBColor(0, 0, 255)
        new_run.font.underline = True

        hyperlink.append(new_run.element)
        paragraph._p.append(hyperlink)

        return hyperlink

    def save(self):
        try:
            self.doc.save(self.path)
        except PermissionError:
            print("Check if the file is open.")
            doc = Document(self.path)
            doc.save(self.path)

    def open(self):
        os.startfile(self.path)

